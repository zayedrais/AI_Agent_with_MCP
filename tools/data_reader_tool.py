import logging
import os
from typing import Dict, Any, Optional, List
import pandas as pd
import docx
import PyPDF2
from pathlib import Path

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger("data-reader-tool")

class DataReaderTool:
    """
    A tool for reading various data file formats including PDF, Excel, and Word documents.
    
    This tool provides functionality to read and extract data from different file types,
    preparing the data for analysis by other tools.
    """
    
    def __init__(self, input_directory: str = None):
        """
        Initialize the data reader tool
        
        Args:
            input_directory: Base directory for input files. If None, use current directory/input_data
        """
        self.base_directory = input_directory or os.path.join(os.getcwd(), "input_data")
        
        # Create input data directory if it doesn't exist
        if not os.path.exists(self.base_directory):
            os.makedirs(self.base_directory)
            logger.info(f"Created input data directory at {self.base_directory}")
    
    def list_available_files(self, file_types: List[str] = None) -> Dict[str, Any]:
        """
        List available files in the input directory, optionally filtered by type
        
        Args:
            file_types: List of file extensions to filter (e.g., ['pdf', 'xlsx'])
            
        Returns:
            Dict containing file information by category
        """
        if file_types is None:
            file_types = ['pdf', 'xlsx', 'xls', 'csv', 'docx', 'doc']
        
        result = {
            "success": True,
            "files": {},
            "count": 0
        }
        
        # Group files by type
        for ext in file_types:
            files = list(Path(self.base_directory).glob(f"**/*.{ext}"))
            if files:
                result["files"][ext] = [str(f.relative_to(self.base_directory)) for f in files]
                result["count"] += len(files)
        
        if result["count"] == 0:
            result["message"] = f"No files found in {self.base_directory} with types: {', '.join(file_types)}"
        else:
            result["message"] = f"Found {result['count']} files in {self.base_directory}"
            
        return result
    
    def read_excel_file(self, file_path: str, sheet_name: Optional[str] = None) -> Dict[str, Any]:
        """
        Read data from Excel file
        
        Args:
            file_path: Path to the Excel file (relative to base_directory)
            sheet_name: Optional sheet name to read (if None, read first sheet)
            
        Returns:
            Dict containing the Excel data and metadata
        """
        try:
            full_path = os.path.join(self.base_directory, file_path)
            
            # Check file exists
            if not os.path.isfile(full_path):
                return {
                    "success": False,
                    "error": f"File not found: {full_path}",
                    "message": "The specified Excel file does not exist"
                }
            
            logger.info(f"Reading Excel file: {full_path}")
            
            # Get sheet names
            excel_file = pd.ExcelFile(full_path)
            sheet_names = excel_file.sheet_names
            
            # If no sheet specified, use the first one
            if sheet_name is None and sheet_names:
                sheet_name = sheet_names[0]
            
            # Read the data
            df = pd.read_excel(full_path, sheet_name=sheet_name)
            
            # Basic data info
            row_count = len(df)
            col_count = len(df.columns)
            data_types = {col: str(dtype) for col, dtype in df.dtypes.items()}
            
            # Convert to dict for serialization (up to 100 rows for preview)
            max_rows = min(100, row_count)
            preview_data = df.head(max_rows).to_dict(orient='records')
            
            return {
                "success": True,
                "file_path": file_path,
                "sheet_name": sheet_name,
                "sheet_names": sheet_names,
                "row_count": row_count,
                "column_count": col_count,
                "columns": list(df.columns),
                "data_types": data_types,
                "preview_data": preview_data,
                "full_data": df,  # Include full DataFrame for analysis
                "message": f"Successfully read Excel file with {row_count} rows and {col_count} columns"
            }
            
        except Exception as e:
            logger.error(f"Error reading Excel file {file_path}: {str(e)}")
            return {
                "success": False,
                "error": str(e),
                "message": f"Failed to read Excel file: {file_path}"
            }
    
    def read_csv_file(self, file_path: str) -> Dict[str, Any]:
        """
        Read data from CSV file
        
        Args:
            file_path: Path to the CSV file (relative to base_directory)
            
        Returns:
            Dict containing the CSV data and metadata
        """
        try:
            full_path = os.path.join(self.base_directory, file_path)
            
            # Check file exists
            if not os.path.isfile(full_path):
                return {
                    "success": False,
                    "error": f"File not found: {full_path}",
                    "message": "The specified CSV file does not exist"
                }
            
            logger.info(f"Reading CSV file: {full_path}")
            
            # Read the data
            df = pd.read_csv(full_path)
            
            # Basic data info
            row_count = len(df)
            col_count = len(df.columns)
            data_types = {col: str(dtype) for col, dtype in df.dtypes.items()}
            
            # Convert to dict for serialization (up to 100 rows for preview)
            max_rows = min(100, row_count)
            preview_data = df.head(max_rows).to_dict(orient='records')
            
            return {
                "success": True,
                "file_path": file_path,
                "row_count": row_count,
                "column_count": col_count,
                "columns": list(df.columns),
                "data_types": data_types,
                "preview_data": preview_data,
                "full_data": df,  # Include full DataFrame for analysis
                "message": f"Successfully read CSV file with {row_count} rows and {col_count} columns"
            }
            
        except Exception as e:
            logger.error(f"Error reading CSV file {file_path}: {str(e)}")
            return {
                "success": False,
                "error": str(e),
                "message": f"Failed to read CSV file: {file_path}"
            }
    
    def read_pdf_file(self, file_path: str) -> Dict[str, Any]:
        """
        Extract text from PDF file
        
        Args:
            file_path: Path to the PDF file (relative to base_directory)
            
        Returns:
            Dict containing the extracted text and metadata
        """
        try:
            full_path = os.path.join(self.base_directory, file_path)
            
            # Check file exists
            if not os.path.isfile(full_path):
                return {
                    "success": False,
                    "error": f"File not found: {full_path}",
                    "message": "The specified PDF file does not exist"
                }
            
            logger.info(f"Reading PDF file: {full_path}")
            
            # Extract text from PDF
            text_content = []
            page_count = 0
            
            with open(full_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                page_count = len(pdf_reader.pages)
                
                for page_num in range(page_count):
                    page = pdf_reader.pages[page_num]
                    text = page.extract_text()
                    text_content.append({
                        "page": page_num + 1,
                        "text": text
                    })
            
            return {
                "success": True,
                "file_path": file_path,
                "page_count": page_count,
                "text_content": text_content,
                "message": f"Successfully extracted text from {page_count} pages of PDF"
            }
            
        except Exception as e:
            logger.error(f"Error reading PDF file {file_path}: {str(e)}")
            return {
                "success": False,
                "error": str(e),
                "message": f"Failed to read PDF file: {file_path}"
            }
    
    def read_word_file(self, file_path: str) -> Dict[str, Any]:
        """
        Extract text from Word document
        
        Args:
            file_path: Path to the Word file (relative to base_directory)
            
        Returns:
            Dict containing the extracted text and metadata
        """
        try:
            full_path = os.path.join(self.base_directory, file_path)
            
            # Check file exists
            if not os.path.isfile(full_path):
                return {
                    "success": False,
                    "error": f"File not found: {full_path}",
                    "message": "The specified Word file does not exist"
                }
            
            logger.info(f"Reading Word file: {full_path}")
            
            # Open the document
            doc = docx.Document(full_path)
            
            # Extract text by paragraphs
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():  # Only include non-empty paragraphs
                    paragraphs.append(para.text)
            
            # Extract tables if any
            tables = []
            for i, table in enumerate(doc.tables):
                table_data = []
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    table_data.append(row_data)
                tables.append({
                    "table_index": i,
                    "data": table_data
                })
            
            return {
                "success": True,
                "file_path": file_path,
                "paragraph_count": len(paragraphs),
                "table_count": len(tables),
                "paragraphs": paragraphs,
                "tables": tables,
                "message": f"Successfully extracted text from Word document with {len(paragraphs)} paragraphs and {len(tables)} tables"
            }
            
        except Exception as e:
            logger.error(f"Error reading Word file {file_path}: {str(e)}")
            return {
                "success": False,
                "error": str(e),
                "message": f"Failed to read Word file: {file_path}"
            }
    
    def read_file(self, file_path: str, sheet_name: Optional[str] = None) -> Dict[str, Any]:
        """
        Read any supported file type based on extension
        
        Args:
            file_path: Path to the file (relative to base_directory)
            sheet_name: Optional sheet name for Excel files
            
        Returns:
            Dict containing the file data and metadata
        """
        logger.info(f"Reading file with DataReaderTool : {file_path}")
        if file_path.lower().endswith(('.xlsx', '.xls')):
            return self.read_excel_file(file_path, sheet_name)
        elif file_path.lower().endswith('.csv'):
            return self.read_csv_file(file_path)
        elif file_path.lower().endswith('.pdf'):
            return self.read_pdf_file(file_path)
        elif file_path.lower().endswith(('.docx', '.doc')):
            return self.read_word_file(file_path)
        else:
            return {
                "success": False,
                "error": f"Unsupported file type: {file_path}",
                "message": "Supported file types are: Excel (.xlsx, .xls), CSV, PDF, and Word (.docx, .doc)"
            }