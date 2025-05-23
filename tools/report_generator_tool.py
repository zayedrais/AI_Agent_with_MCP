import logging
import os
from typing import Dict, Any, List, Optional
from datetime import datetime
import base64
import io
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import RGBColor
from pathlib import Path

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger("report-generator-tool")

class ReportGeneratorTool:
    """
    A tool for generating formatted reports with data analysis results and visualizations
    
    This tool can create Word documents with text, tables, and visualizations from
    data analysis results.
    """
    
    def __init__(self, output_directory: str = None):
        """
        Initialize the report generator tool
        
        Args:
            output_directory: Directory to save reports. If None, use current directory/reports
        """
        logger.info("Initializing ReportGeneratorTool")
        self.output_directory = output_directory or os.path.join(os.getcwd(), "reports")
        logger.debug(f"Setting output directory to: {self.output_directory}")
        
        # Create output directory if it doesn't exist
        if not os.path.exists(self.output_directory):
            os.makedirs(self.output_directory)
            logger.info(f"Created reports directory at {self.output_directory}")
    
    def create_report(self, 
                     title: str, 
                     sections: List[Dict[str, Any]],
                     file_name: str = None,
                     author: str = "Data Analysis Agent",
                     include_timestamp: bool = True) -> Dict[str, Any]:
        """
        Create a Word document report with multiple sections
        
        Args:
            title: The title of the report
            sections: A list of sections, each with title, content, and optionally tables or images
            file_name: Name for the output file (without path). If None, generated from title
            author: The author of the report
            include_timestamp: Whether to include a timestamp in the file name
            
        Returns:
            Dict containing the result of the operation
        """
        try:
            logger.info(f"Creating report with title: '{title}'")
            logger.debug(f"Report parameters - sections: {len(sections)}, author: {author}, include_timestamp: {include_timestamp}")
            
            # Create a new Document
            doc = Document()
            logger.debug("Created new Document object")
            
            # Set up document properties
            doc.core_properties.title = title
            doc.core_properties.author = author
            doc.core_properties.created = datetime.now()
            logger.debug(f"Set document properties - title: '{title}', author: '{author}'")
            
            # Set up document styles
            self._set_up_document_styles(doc)
            logger.debug("Applied document styles")
            
            # Add a title
            doc.add_heading(title, level=0)
            logger.debug("Added document title heading")
            
            # Add timestamp if needed
            if include_timestamp:
                timestamp_para = doc.add_paragraph()
                timestamp_para.add_run(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}").italic = True
                timestamp_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                logger.debug("Added timestamp to document")
            
            # Add sections
            logger.info(f"Adding {len(sections)} sections to document")
            for i, section in enumerate(sections):
                logger.debug(f"Processing section {i+1}/{len(sections)}: '{section.get('title', 'Untitled')}'")
                self._add_section_to_document(doc, section)
            
            # Generate file name if not provided
            if not file_name:
                clean_title = "".join(c if c.isalnum() else "_" for c in title)
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                file_name = f"{clean_title}_{timestamp}.docx"
                logger.debug(f"Generated file name: {file_name}")
            
            # Ensure file has .docx extension
            if not file_name.lower().endswith('.docx'):
                file_name += '.docx'
                logger.debug(f"Added .docx extension to file name: {file_name}")
            
            # Save the document
            file_path = os.path.join(self.output_directory, file_name)
            logger.debug(f"Saving document to: {file_path}")
            doc.save(file_path)
            logger.info(f"Successfully saved report to {file_path}")
            
            return {
                "success": True,
                "file_path": file_path,
                "file_name": file_name,
                "section_count": len(sections),
                "message": f"Successfully generated report at {file_path}"
            }
            
        except Exception as e:
            logger.error(f"Error creating report: {str(e)}", exc_info=True)
            return {
                "success": False,
                "error": str(e),
                "message": "Failed to create report"
            }
    
    def _set_up_document_styles(self, doc: Document):
        """
        Set up document styles for consistent formatting
        
        Args:
            doc: The document object to set up styles for
        """
        logger.debug("Setting up document styles")
        
        # Heading styles
        for i in range(1, 5):
            style = doc.styles[f'Heading {i}']
            font = style.font
            font.name = 'Calibri'
            font.size = Pt(20 - (i * 2))  # Decrease size for each heading level
            font.bold = True
            if i == 1:
                font.color.rgb = RGBColor(0x35, 0x35, 0x95)  # Dark blue for Heading 1
            logger.debug(f"Configured Heading {i} style - font: {font.name}, size: {font.size.pt}pt")
        
        # Normal text style
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)
        logger.debug(f"Configured Normal style - font: {font.name}, size: {font.size.pt}pt")
        
        # Create a style for tables
        try:
            table_style = doc.styles.add_style('TableStyle', WD_STYLE_TYPE.PARAGRAPH)
            font = table_style.font
            font.name = 'Calibri'
            font.size = Pt(10)
            logger.debug("Created TableStyle")
        except Exception as e:
            logger.debug(f"TableStyle already exists: {str(e)}")
        
        # Create a caption style
        try:
            caption_style = doc.styles.add_style('Caption', WD_STYLE_TYPE.PARAGRAPH)
            font = caption_style.font
            font.name = 'Calibri'
            font.size = Pt(10)
            font.italic = True
            logger.debug("Created Caption style")
        except Exception as e:
            logger.debug(f"Caption style already exists: {str(e)}")
    
    def _add_section_to_document(self, doc: Document, section: Dict[str, Any]):
        """
        Add a section to the document
        
        Args:
            doc: The document object to add the section to
            section: The section data dictionary
        """
        # Section must have title and content
        title = section.get("title", "")
        content = section.get("content", "")
        level = section.get("level", 1)  # Heading level
        
        logger.debug(f"Adding section: '{title}' (level {level})")
        
        # Add section heading
        if title:
            heading = doc.add_heading(title, level=level)
            # Add page break if specified
            if section.get("page_break_before", False):
                heading.paragraph_format.page_break_before = True
                logger.debug(f"Added page break before section '{title}'")
        
        # Add content paragraphs
        if content:
            # Check if content is a string or a list
            if isinstance(content, str):
                content_paras = content.split("\n\n")  # Split on double newlines
                logger.debug(f"Split content into {len(content_paras)} paragraphs")
            else:
                content_paras = content
                logger.debug(f"Using pre-split content with {len(content_paras)} paragraphs")
            
            # Add each paragraph
            for i, para_text in enumerate(content_paras):
                if para_text.strip():
                    para = doc.add_paragraph(para_text.strip())
                    logger.debug(f"Added paragraph {i+1}/{len(content_paras)} ({len(para_text)} chars)")
        
        # Add tables if present
        tables = section.get("tables", [])
        if tables:
            logger.debug(f"Adding {len(tables)} tables to section '{title}'")
            
        for i, table_data in enumerate(tables):
            table_title = table_data.get("title", "")
            headers = table_data.get("headers", [])
            rows = table_data.get("rows", [])
            
            logger.debug(f"Processing table {i+1}/{len(tables)}: '{table_title}' with {len(rows)} rows")
            
            if headers and rows:
                # Add a table title/caption
                if table_title:
                    caption = doc.add_paragraph(f"Table: {table_title}")
                    caption.style = doc.styles["Caption"]
                    logger.debug(f"Added table caption: {table_title}")
                
                # Create the table
                table = doc.add_table(rows=1, cols=len(headers))
                table.style = 'Table Grid'
                logger.debug(f"Created table with {len(headers)} columns")
                
                # Add headers
                header_cells = table.rows[0].cells
                for i, header in enumerate(headers):
                    header_cells[i].text = str(header)
                    # Make headers bold
                    for paragraph in header_cells[i].paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
                logger.debug(f"Added table headers: {headers}")
                
                # Add rows
                for j, row_data in enumerate(rows):
                    row_cells = table.add_row().cells
                    for i, cell_data in enumerate(row_data):
                        if i < len(row_cells):  # Prevent index errors
                            row_cells[i].text = str(cell_data)
                logger.debug(f"Added {len(rows)} rows to table")
                
                # Add some space after the table
                doc.add_paragraph()
        
        # Add images if present
        images = section.get("images", [])
        if images:
            logger.debug(f"Adding {len(images)} images to section '{title}'")
            
        for i, img_data in enumerate(images):
            img_title = img_data.get("title", "")
            img_base64 = img_data.get("image_base64", "")
            width_inches = img_data.get("width_inches", 6.0)
            description = img_data.get("description", "")
            
            logger.debug(f"Processing image {i+1}/{len(images)}: '{img_title}' (width: {width_inches} inches)")
            
            if img_base64:
                # Add image title
                if img_title:
                    caption = doc.add_paragraph(f"Figure: {img_title}")
                    caption.style = doc.styles["Caption"]
                    logger.debug(f"Added image caption: {img_title}")
                
                try:
                    # Add the image
                    image_stream = io.BytesIO(base64.b64decode(img_base64))
                    doc.add_picture(image_stream, width=Inches(width_inches))
                    logger.debug(f"Added image from base64 data (decoded length: {len(base64.b64decode(img_base64))})")
                    
                    # Add image description if provided
                    if description:
                        desc_para = doc.add_paragraph(description)
                        desc_para.style = 'Caption'
                        desc_para.italic = True
                        logger.debug(f"Added image description: {description[:50]}...")
                    
                    # Add some space after the image
                    doc.add_paragraph()
                except Exception as e:
                    logger.error(f"Failed to add image {i+1}: {str(e)}")
    
    def generate_report_template(self, title: str, file_name: str = None) -> Dict[str, Any]:
        """
        Generate an empty report template with standard sections
        
        Args:
            title: The title of the report
            file_name: Name for the output file (without path). If None, generated from title
            
        Returns:
            Dict containing the result of the operation and template structure
        """
        try:
            logger.info(f"Generating report template with title with Tool: '{title}'")
            
            # Define standard sections
            sections = [
                {
                    "title": "Executive Summary",
                    "content": "Brief overview of the analysis and key findings.",
                    "level": 1
                },
                {
                    "title": "Introduction",
                    "content": "Context and objectives of the data analysis.",
                    "level": 1
                },
                {
                    "title": "Data Overview",
                    "content": "Description of the data sources and structure.",
                    "level": 1,
                    "tables": [
                        {
                            "title": "Dataset Summary",
                            "headers": ["Attribute", "Value"],
                            "rows": [
                                ["Data Source", ""],
                                ["Number of Records", ""],
                                ["Time Period", ""],
                                ["Key Variables", ""]
                            ]
                        }
                    ]
                },
                {
                    "title": "Methodology",
                    "content": "Methods and techniques used in the analysis.",
                    "level": 1
                },
                {
                    "title": "Key Findings",
                    "content": "Main results and insights from the analysis.",
                    "level": 1,
                    "page_break_before": True
                },
                {
                    "title": "Data Visualizations",
                    "content": "Visual representation of key insights.",
                    "level": 1
                },
                {
                    "title": "Recommendations",
                    "content": "Actionable insights based on the analysis.",
                    "level": 1
                },
                {
                    "title": "Conclusion",
                    "content": "Summary of findings and next steps.",
                    "level": 1
                },
                {
                    "title": "Appendix",
                    "content": "Additional details and supporting information.",
                    "level": 1,
                    "page_break_before": True
                }
            ]
            
            logger.debug(f"Created template with {len(sections)} standard sections")
            
            # Create the report
            result = self.create_report(
                title=title,
                sections=sections,
                file_name=file_name
            )
            
            # If successful, add template structure to result
            if result.get("success", False):
                result["template_structure"] = [section["title"] for section in sections]
                result["message"] = f"Successfully generated report template at {result['file_path']}"
                logger.info(f"Successfully generated report template at {result.get('file_path')}")
            
            return result
            
        except Exception as e:
            logger.error(f"Error generating report template: {str(e)}", exc_info=True)
            return {
                "success": False,
                "error": str(e),
                "message": "Failed to generate report template"
            }

    def create_markdown_report(self, 
                         title: str,
                         sections: List[Dict[str, Any]],
                         file_name: str = None,
                         author: str = "Data Analysis Agent",
                         include_timestamp: bool = True) -> Dict[str, Any]:
        """
        Create a Markdown (.md) format report with multiple sections
        
        Args:
            title: The title of the report
            sections: A list of sections, each with title, content, and optionally tables or images
            file_name: Name for the output file (without path). If None, generated from title
            author: The author of the report
            include_timestamp: Whether to include a timestamp in the file name and content
            
        Returns:
            Dict containing the result of the operation
        """
        try:
            logger.info(f"Creating markdown report with title: '{title}'")
            logger.debug(f"Report parameters - sections: {len(sections)}, author: {author}, include_timestamp: {include_timestamp}")
            
            # Start building the markdown content
            markdown_content = f"# {title}\n\n"
            
            # Add metadata
            if author:
                markdown_content += f"**Author:** {author}  \n"
            if include_timestamp:
                timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
                markdown_content += f"**Generated:** {timestamp}  \n"
            
            markdown_content += "\n"
            
            # Process each section
            logger.info(f"Adding {len(sections)} sections to markdown document")
            for i, section in enumerate(sections):
                logger.debug(f"Processing section {i+1}/{len(sections)}: '{section.get('title', 'Untitled')}'")
                
                # Add section title with appropriate heading level
                title = section.get("title", "")
                level = section.get("level", 1)
                
                if title:
                    # Add page break using HTML if specified (though not all MD renderers will respect this)
                    if section.get("page_break_before", False):
                        markdown_content += "<div style='page-break-before: always'></div>\n\n"
                    
                    # Add section heading with proper level (# to ######)
                    heading_markers = "#" * min(level + 1, 6)  # +1 because title is h1
                    markdown_content += f"{heading_markers} {title}\n\n"
                
                # Add content
                content = section.get("content", "")
                if content:
                    if isinstance(content, str):
                        markdown_content += f"{content}\n\n"
                    elif isinstance(content, list):
                        for para in content:
                            markdown_content += f"{para}\n\n"
                
                # Add tables if present
                tables = section.get("tables", [])
                for j, table_data in enumerate(tables):
                    table_title = table_data.get("title", "")
                    headers = table_data.get("headers", [])
                    rows = table_data.get("rows", [])
                    
                    logger.debug(f"Processing table {j+1}/{len(tables)}: '{table_title}' with {len(rows)} rows")
                    
                    if table_title:
                        markdown_content += f"**Table: {table_title}**\n\n"
                    
                    if headers and rows:
                        # Create markdown table headers
                        markdown_content += "| " + " | ".join(str(h) for h in headers) + " |\n"
                        markdown_content += "| " + " | ".join(["---"] * len(headers)) + " |\n"
                        
                        # Add rows
                        for row_data in rows:
                            markdown_content += "| " + " | ".join(str(cell) for cell in row_data) + " |\n"
                        
                        markdown_content += "\n"
                
                # Add images if present
                images = section.get("images", [])
                for j, img_data in enumerate(images):
                    img_title = img_data.get("title", "")
                    img_base64 = img_data.get("image_base64", "")
                    description = img_data.get("description", "")
                    
                    logger.debug(f"Processing image {j+1}/{len(images)}: '{img_title}'")
                    
                    if img_base64:
                        try:
                            # Save image to file
                            img_file_name = f"image_{i}_{j}.png"
                            img_dir = os.path.join(self.output_directory, "images")
                            os.makedirs(img_dir, exist_ok=True)
                            
                            img_path = os.path.join(img_dir, img_file_name)
                            with open(img_path, "wb") as f:
                                f.write(base64.b64decode(img_base64))
                            
                            # Reference the image in markdown using relative path
                            if img_title:
                                markdown_content += f"**Figure: {img_title}**\n\n"
                            
                            # Include image with reference to the saved file (relative path)
                            markdown_content += f"![{img_title or 'Image'}](images/{img_file_name})\n\n"
                            
                            if description:
                                markdown_content += f"*{description}*\n\n"
                                
                        except Exception as e:
                            logger.error(f"Failed to add image {j+1}: {str(e)}")
                            # Add a note about the error
                            markdown_content += f"*[Image could not be included: {str(e)}]*\n\n"
            
            # Generate file name if not provided
            if not file_name:
                clean_title = "".join(c if c.isalnum() else "_" for c in title)
                timestamp_str = datetime.now().strftime("%Y%m%d_%H%M%S")
                file_name = f"{clean_title}_{timestamp_str}.md"
                logger.debug(f"Generated file name: {file_name}")
            
            # Ensure file has .md extension
            if not file_name.lower().endswith('.md'):
                file_name += '.md'
                logger.debug(f"Added .md extension to file name: {file_name}")
            
            # Save the markdown content to file
            file_path = os.path.join(self.output_directory, file_name)
            logger.debug(f"Saving markdown document to: {file_path}")
            
            with open(file_path, "w", encoding="utf-8") as f:
                f.write(markdown_content)
            
            logger.info(f"Successfully saved markdown report to {file_path}")
            
            return {
                "success": True,
                "file_path": file_path,
                "file_name": file_name,
                "section_count": len(sections),
                "message": f"Successfully generated markdown report at {file_path}",
                "format": "markdown"
            }
            
        except Exception as e:
            logger.error(f"Error creating markdown report: {str(e)}", exc_info=True)
            return {
                "success": False,
                "error": str(e),
                "message": "Failed to create markdown report"
            }